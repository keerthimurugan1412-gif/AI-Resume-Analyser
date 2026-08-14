"""
Resume Upload and Text Extraction:

Responsible for:
- Validating uploaded resume files (type / size).
- Extracting raw text from PDF and DOCX resumes.
- Never persisting the uploaded file to disk unless the caller
  explicitly opts in (Responsible-AI rule: "Avoid storing resumes
  permanently unless the user gives permission").
"""

from __future__ import annotations

import io
from dataclasses import dataclass
from typing import Optional, Union

from pypdf import PdfReader
from docx import Document

ALLOWED_EXTENSIONS = (".pdf", ".docx")
MAX_FILE_SIZE_MB = 5


@dataclass
class ExtractionResult:
    """Container for the outcome of a resume extraction attempt."""
    success: bool
    text: str = ""
    file_name: str = ""
    file_type: str = ""
    error: Optional[str] = None


def validate_file(file_name: str, file_size_bytes: int) -> Optional[str]:
    """
    Validate file type and size.

    Returns None if the file is valid, otherwise an error message.
    """
    lower_name = file_name.lower()
    if not lower_name.endswith(ALLOWED_EXTENSIONS):
        return f"Unsupported file type. Please upload a PDF or DOCX file (got '{file_name}')."

    max_bytes = MAX_FILE_SIZE_MB * 1024 * 1024
    if file_size_bytes > max_bytes:
        return f"File is too large ({file_size_bytes / (1024 * 1024):.1f} MB). Max allowed size is {MAX_FILE_SIZE_MB} MB."

    if file_size_bytes == 0:
        return "The uploaded file appears to be empty."

    return None


def _extract_pdf_text(file_bytes: bytes) -> str:
    """Extract text from every page of a PDF resume."""
    reader = PdfReader(io.BytesIO(file_bytes))
    pages_text = []
    for page in reader.pages:
        page_text = page.extract_text() or ""
        pages_text.append(page_text)
    return "\n".join(pages_text)


def _extract_docx_text(file_bytes: bytes) -> str:
    """Extract text from every paragraph (and table cell) of a DOCX resume."""
    document = Document(io.BytesIO(file_bytes))
    parts = [p.text for p in document.paragraphs]

    # Many resumes place skills / experience inside tables.
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)

    return "\n".join(parts)


def extract_resume_text(file_name: str, file_bytes: bytes) -> ExtractionResult:
    """
    Main entry point: validates and extracts text from an uploaded resume.

    Parameters
    ----------
    file_name: original uploaded file name (used to detect type).
    file_bytes: raw bytes of the uploaded file (kept in-memory only).
    """
    error = validate_file(file_name, len(file_bytes))
    if error:
        return ExtractionResult(success=False, file_name=file_name, error=error)

    lower_name = file_name.lower()
    try:
        if lower_name.endswith(".pdf"):
            text = _extract_pdf_text(file_bytes)
            file_type = "pdf"
        else:
            text = _extract_docx_text(file_bytes)
            file_type = "docx"
    except Exception as exc:  # noqa: BLE001 - surface a friendly error to the UI
        return ExtractionResult(
            success=False,
            file_name=file_name,
            error=f"Could not read the file. It may be corrupted or password protected. ({exc})",
        )

    if not text.strip():
        return ExtractionResult(
            success=False,
            file_name=file_name,
            file_type=file_type,
            error="No readable text was found in this file (it may be a scanned/image-only resume).",
        )

    return ExtractionResult(success=True, text=text, file_name=file_name, file_type=file_type)


def extract_from_path(path: str) -> ExtractionResult:
    """Convenience helper used by tests / CLI runs to read a resume from disk."""
    with open(path, "rb") as f:
        data = f.read()
    file_name = path.split("/")[-1]
    return extract_resume_text(file_name, data)
